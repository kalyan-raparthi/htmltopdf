from flask import Flask, request, send_file
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches
import os
import imghdr
import tempfile

app = Flask(__name__)

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    data = request.get_json()
    url = data.get('url')

    if not url or not url.startswith('http'):
        return {'error': 'Invalid URL'}, 400

    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, 'output.docx')
    image_dir = os.path.join(temp_dir, 'images')
    os.makedirs(image_dir, exist_ok=True)


    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')

        text_content = soup.get_text(separator='\n', strip=True)

        images = []
        for idx, img in enumerate(soup.find_all('img'), 1):
            src = img.get('src')
            if src:
                full_url = urljoin(url, src)
                try:
                    img_data = requests.get(full_url, timeout=10).content
                    img_file = os.path.join(image_dir, f'image_{idx}')
                    with open(img_file, 'wb') as f:
                        f.write(img_data)

                    ext = imghdr.what(img_file)
                    if ext in ['jpeg', 'png', 'bmp', 'gif']:
                        valid_file = img_file + f'.{ext}'
                        os.rename(img_file, valid_file)
                        images.append(valid_file)
                    else:
                        os.remove(img_file)
                except:
                    continue
    except Exception as e:
        return {'error': f'Failed to fetch page: {str(e)}'}, 500

    doc = Document()
    doc.add_heading('Web Page Content', 0)
    doc.add_paragraph(text_content)

    doc.add_page_break()
    doc.add_heading('Images', level=1)
    doc.add_paragraph(f'Page URL: {url}')

    for img in images:
        try:
            doc.add_picture(img, width=Inches(5))
        except:
            continue

    doc.save(docx_path)

    return send_file(docx_path, as_attachment=True, download_name='page_output.docx')


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
